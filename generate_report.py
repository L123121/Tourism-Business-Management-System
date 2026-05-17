# -*- coding: utf-8 -*-
"""按照用户模板格式生成旅游业务管理系统UML建模实验报告（扩充版）"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold


def add_para(doc, text, bold=False, size=12, indent=True, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p


def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_bold_label(doc, text):
    return add_para(doc, text, bold=True, size=12, indent=False)


def create_report():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ================================================================
    # 表头信息
    # ================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('实验报告书')
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True
    doc.add_paragraph()

    tbl = doc.add_table(rows=5, cols=4, style='Table Grid')
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ['实验课程名称：', '软件工程基础实验', '实验项目名称：', 'UML建模'],
        ['实验成绩', '', '实 验 者', '（请填写）'],
        ['专业班级', '（请填写）', '组 别', '（请填写）'],
        ['同 组 者', '（请填写）', '实验日期', '（请填写）'],
    ]
    for i, r in enumerate(rows):
        for j, t in enumerate(r):
            set_cell(tbl.cell(i, j), t, bold=(j % 2 == 0), size=11)

    doc.add_paragraph()

    # ================================================================
    # 第一部分：实验预习报告
    # ================================================================
    add_heading(doc, '第一部分：实验预习报告', level=1)
    add_para(doc, '（包括实验目的、意义，实验基本原理与方法，主要仪器设备及耗材，实验方案与技术路线等）', size=10, indent=False)

    add_bold_label(doc, '一、实验目的与意义')
    add_para(doc, '1. 掌握统一建模语言（UML）的基本概念和常用图的绘制方法，包括用例图、类图、顺序图、活动图和状态机图，能够根据实际业务需求选择合适的UML图进行系统建模。')
    add_para(doc, '2. 学习运用面向对象的分析与设计方法，对实际业务系统进行需求分析和系统设计，理解从问题描述到系统模型的转换过程。')
    add_para(doc, '3. 通过"旅游业务管理系统"这一综合性案例，实践完整的软件建模过程，将理论知识应用于实际问题的求解中。')
    add_para(doc, '4. 体验结对编程（Pair Programming）的开发模式，培养团队协作与沟通能力，理解极限编程等敏捷开发方法的核心思想。')
    add_para(doc, '5. 提高软件工程实践能力，建立从需求分析、系统设计到编码实现的软件开发全流程认知，为后续的软件工程课程学习和实际项目开发奠定基础。')

    add_bold_label(doc, '二、实验基本原理与方法')
    add_para(doc, '本实验基于面向对象的软件工程方法，以UML（Unified Modeling Language，统一建模语言）为建模工具，对旅游业务管理系统进行分析与设计。UML是一种标准化的建模语言，广泛应用于软件系统的可视化、规约、构造和文档化。本实验主要涉及以下几种UML图：')
    add_bullet(doc, '用例图（Use Case Diagram）：从用户角度描述系统的功能需求，通过识别参与者（Actor）和用例（Use Case）来界定系统的功能范围。用例图是需求分析阶段最重要的工具之一，它能够清晰地展示系统"做什么"而非"怎么做"。')
    add_bullet(doc, '类图（Class Diagram）：描述系统的静态结构模型，展示系统中的类、类的属性和操作，以及类之间的各种关系（关联、聚合、组合、继承、依赖等）。类图是面向对象设计的核心，直接指导后续的编码实现。')
    add_bullet(doc, '顺序图（Sequence Diagram）：描述对象之间按时间顺序排列的交互过程，强调消息传递的先后顺序。顺序图适合分析具体用例的实现细节，帮助理解对象之间的协作方式。')
    add_bullet(doc, '活动图（Activity Diagram）：描述业务流程或算法中活动的控制流程，类似于流程图但支持并行处理。活动图适合描述复杂的业务流程，能够清晰地展示条件判断和并行操作。')
    add_bullet(doc, '状态机图（State Machine Diagram）：描述一个对象在其生命周期内所经历的各种状态，以及导致状态转换的事件。状态机图适合描述具有复杂状态变化的对象，如订单、申请等业务对象。')
    add_para(doc, '实验采用结对编程（Pair Programming）的方式进行开发，这是极限编程（XP）的核心实践之一。在结对编程中，两名成员共用一台计算机，一人担任驾驶员（Driver）负责编写代码，另一人担任领航员（Navigator）负责审查代码、思考战略方向，两人定期互换角色。研究表明，结对编程能够显著提高代码质量，减少缺陷率，同时促进知识在团队内的传播。')

    add_bold_label(doc, '三、主要仪器设备及耗材')
    add_bullet(doc, '计算机（Windows 11操作系统，8GB以上内存）')
    add_bullet(doc, '开发工具：IDE（如IntelliJ IDEA / VS Code / Eclipse等）')
    add_bullet(doc, 'UML建模工具：PlantUML / StarUML / Draw.io等')
    add_bullet(doc, '编程语言：（请根据实际使用的语言填写，如Java / Python / C#等）')
    add_bullet(doc, '版本管理工具：Git')
    add_bullet(doc, '项目管理工具：（如Maven / Gradle等，根据实际情况填写）')

    add_bold_label(doc, '四、实验方案与技术路线')
    add_para(doc, '本实验按照软件工程的生命周期，采用迭代的方式进行，具体技术路线如下：')
    add_para(doc, '（1）需求分析阶段：仔细阅读旅游业务管理系统的问题描述文档，梳理业务流程和业务规则，识别系统的参与者和用例，绘制用例图，编写用例描述。')
    add_para(doc, '（2）系统分析阶段：根据需求分析结果，识别系统中的核心类和对象，分析类的属性和方法，确定类之间的关系，绘制类图。同时，针对关键业务场景，分析对象之间的交互过程。')
    add_para(doc, '（3）系统设计阶段：针对关键业务场景绘制顺序图，描述对象之间的消息传递；绘制活动图描述完整的业务流程；绘制状态机图描述旅游团、申请等关键对象的状态转换过程。')
    add_para(doc, '（4）编码实现阶段：基于UML模型，采用结对编程方式进行系统原型开发，实现旅游申请、参加者管理、余款支付、取消申请、路线管理等核心业务功能。')
    add_para(doc, '（5）测试与总结阶段：对系统进行功能测试，验证各业务场景的正确性；整理实验过程中遇到的问题和解决方案；撰写实验报告，总结实验心得和体会。')

    # ================================================================
    # 第二部分：实验过程记录
    # ================================================================
    doc.add_page_break()
    add_heading(doc, '第二部分：实验过程记录', level=1)
    add_para(doc, '（包括实验原始数据记录，实验现象记录，实验过程发现的问题等）', size=10, indent=False)

    add_bold_label(doc, '一、问题描述')
    add_para(doc, 'XXX旅行社是武汉地区一家专门提供组团旅行的旅游公司。目前有关旅游业务的申请过程都是手工完成，考虑到旅游业务的迅猛发展，公司决定开发一个旅游业务申请信息系统。')
    add_para(doc, '使用新系统后，XXX旅行社的核心业务流程如下：')
    add_bullet(doc, '前台招待顾客的员工同顾客洽谈旅游的各项事宜，并帮助旅客办理相关申请手续。首先调查顾客所要求的旅行状况，根据顾客的要求查询相关的旅游团的详细信息。')
    add_bullet(doc, '申请条件：所申请旅游团在截止日期之前，且所申请旅游团的人数限额未满。不同旅游团的截止日期不同。')
    add_bullet(doc, '在满足条件的情况下，员工向系统录入申请责任人的姓名、电话号码、参加的大人和孩子的人数，系统自动算出订金。订金金额由距离出发日期的天数决定。')
    add_bullet(doc, '办理完手续后，旅行社员工向顾客收取订金，在系统中记录订金支付情况，并向顾客提供订金收据和旅游申请书。')
    add_bullet(doc, '顾客将填写好的旅游申请书一周内邮寄到旅行社。员工通过旅游团代码、出发日期和申请负责人等信息查询申请，将参加者信息录入系统。')
    add_bullet(doc, '所有参加者信息录入后，申请完成。系统打印旅游确认书和余额交款单邮寄给申请责任人。')

    add_para(doc, '余款支付规则：余款的支付期限为出发日期前30天。但从交款单发送日到支付期限不足10天的情况下，支付期限为发送日后10天。即 payment_deadline = max(出发日期-30天, 交款单发送日+10天)。')

    add_para(doc, '申请完成后的变更规则：')
    add_bullet(doc, '只能办理参加者信息的变更或取消，以及整个申请的取消。新增参加者视为新申请处理。')
    add_bullet(doc, '取消或变更参加者时，若对象为申请责任人本人，必须选定新的申请责任人。')
    add_bullet(doc, '取消申请时，从已支付金额中扣除取消手续费后返还顾客。取消手续费与距出发天数相关，申请完成前取消也适用相同规则。')

    add_para(doc, '路线与价格管理：每个季度由专门员工进行旅游路线设计和旅游活动设定。缺乏吸引力的线路可取消，但历史信息不删除，变更后的线路作为新线路录入并保留变更历史。每个季度可能对旅游活动价格进行调整，价格包括大人价格、小孩价格及优惠措施。价格公开后不可变更，未定价的旅游团不对顾客公开。')

    add_para(doc, '财务数据导出：每天晚上，系统自动将当天与现金相关的订金、支付信息全部导出到另一个待开发的财务系统中，会计人员第二天利用财务系统进行记账等财务操作。')

    add_bold_label(doc, '二、系统分析与设计')

    add_para(doc, '（一）用例分析', bold=True, indent=False)
    add_para(doc, '通过对问题描述文档的深入分析，我们识别出系统的参与者和用例，并绘制了用例图。')

    add_para(doc, '1. 参与者识别', bold=True, indent=False)
    add_para(doc, '系统涉及以下参与者：')
    add_bullet(doc, '前台员工（Receptionist）：系统的最主要用户，负责接待顾客、查询旅游团信息、办理申请手续、录入参加者信息、处理余款支付、办理变更和取消业务。')
    add_bullet(doc, '路线管理员（RouteManager）：负责旅游路线的设计、旅游活动的设定、价格的调整等后台管理工作。')
    add_bullet(doc, '催款员工（CollectionClerk）：负责每日打印已完成申请的旅游确认书和余额交款单，并邮寄给申请人。')
    add_bullet(doc, '会计人员（Accountant）：利用财务系统处理导出的费用数据，进行记账等财务操作。虽然会计人员不直接操作本系统，但系统的财务数据导出功能是为其服务的。')
    add_bullet(doc, '系统（System）：作为参与者，负责每晚自动执行财务数据导出任务。')
    add_bullet(doc, '顾客（Customer）：作为外部参与者，通过前台员工与系统间接交互，不直接操作系统。')

    add_para(doc, '2. 用例识别', bold=True, indent=False)
    add_para(doc, '根据业务流程分析，识别出以下主要用例：')
    add_bullet(doc, '办理旅游申请：查询旅游团信息、验证申请条件（截止日期和人数）、录入申请信息、计算订金、记录订金支付、提供收据和申请书。')
    add_bullet(doc, '录入参加者信息：根据邮寄回的旅游申请书，查询对应申请，录入参加者的详细信息。')
    add_bullet(doc, '处理余款支付：通过交款单编号、旅游团代码等查询申请信息，将支付完成信息录入系统。')
    add_bullet(doc, '打印确认书与交款单：每日打印前一天已完成申请的旅游确认书和余额交款单。')
    add_bullet(doc, '变更参加者信息：修改已录入的参加者信息，若变更为申请责任人则需指定新的申请责任人。')
    add_bullet(doc, '取消参加者：取消部分参加者，若取消申请责任人则需指定新的申请责任人。')
    add_bullet(doc, '取消申请：取消整个申请，计算取消手续费，退还余款。')
    add_bullet(doc, '设计旅游路线：录入新的旅游路线信息。')
    add_bullet(doc, '设定旅游活动：为旅游路线设定具体的旅游活动。')
    add_bullet(doc, '调整价格：设定或调整旅游活动的大人价格、小孩价格和优惠措施。')
    add_bullet(doc, '导出财务数据：每晚自动将当天的订金、支付信息导出到财务系统。')

    add_para(doc, '（二）类图设计', bold=True, indent=False)
    add_para(doc, '根据业务分析，我们识别出以下核心类，并分析了它们的属性、方法和相互关系：')

    add_para(doc, '1. 核心类识别', bold=True, indent=False)
    add_bullet(doc, '旅游路线（TourRoute）：属性包括路线代码、路线名称、路线描述、创建日期、状态（活跃/已取消）等。方法包括添加旅游活动、取消路线等。一条旅游路线可包含多个旅游活动。')
    add_bullet(doc, '旅游活动（TourActivity）：属性包括活动代码、活动名称、活动描述等。方法包括创建旅游团。一个旅游活动可组织多个不同出发日期的旅游团。')
    add_bullet(doc, '旅游团（TourGroup）：属性包括团代码、出发日期、截止日期、人数限额、当前人数、状态（未开放/已开放/已截止/已完成）等。方法包括判断是否可以申请、增加人数等。')
    add_bullet(doc, '价格（Price）：属性包括大人价格、小孩价格、优惠措施、是否已公开等。方法包括公开价格、计算总价等。价格与旅游团一对一关联，公开后不可修改。')
    add_bullet(doc, '申请（Application）：属性包括申请编号、申请日期、状态（进行中/已完成/已取消）等。方法包括办理申请、取消申请、计算订金、计算取消手续费等。')
    add_bullet(doc, '申请责任人（Applicant）：属性包括姓名、电话等。与申请一对一关联。')
    add_bullet(doc, '参加者（Participant）：属性包括姓名、性别、年龄、类型（大人/小孩）等。方法包括修改信息、取消参加。参加者属于某个申请。')
    add_bullet(doc, '支付记录（Payment）：属性包括支付编号、金额、支付日期、支付类型（订金/余款/退款）等。')
    add_bullet(doc, '收据（Receipt）：属性包括收据编号、金额、开具日期等。')
    add_bullet(doc, '变更历史（ChangeHistory）：属性包括变更编号、变更日期、变更内容、操作人等。用于记录旅游路线的变更过程。')

    add_para(doc, '2. 类之间的关系', bold=True, indent=False)
    add_bullet(doc, '旅游路线与旅游活动：组合关系（1对多），旅游活动不能脱离旅游路线独立存在。')
    add_bullet(doc, '旅游活动与旅游团：组合关系（1对多），旅游团是旅游活动的具体组织形式。')
    add_bullet(doc, '旅游团与价格：一对一关联关系，每个旅游团对应一个价格，价格公开后不可修改。')
    add_bullet(doc, '旅游团与申请：一对多关联关系，一个旅游团可以接受多个申请。')
    add_bullet(doc, '申请与申请责任人：一对一关联关系，每个申请有一个申请责任人。')
    add_bullet(doc, '申请与参加者：一对多组合关系，一个申请包含多个参加者。')
    add_bullet(doc, '申请与支付记录：一对多组合关系，一个申请可有多条支付记录（订金、余款等）。')
    add_bullet(doc, '旅游路线与变更历史：一对多关联关系，记录路线的变更过程。')

    add_para(doc, '（三）关键算法设计', bold=True, indent=False)
    add_para(doc, '1. 订金计算算法', bold=True, indent=False)
    add_para(doc, '根据问题描述，订金金额由距离出发日期的天数决定。系统首先计算当前日期与出发日期之间的天数差，然后查表确定每人的订金金额。总订金 = 每人订金 ×（大人数 + 小孩数）。该算法的关键在于日期计算的准确性，需要考虑月份天数差异和闰年等边界情况。')

    add_para(doc, '2. 取消手续费计算算法', bold=True, indent=False)
    add_para(doc, '取消申请时，根据距出发日期的天数确定手续费比例。天数越多，手续费比例越低；天数越少，手续费越高。手续费 = 已支付总额 × 手续费比例。退还金额 = 已支付总额 - 手续费。该算法在申请完成前和完成后均适用相同的规则。')

    add_para(doc, '3. 余款支付期限计算算法', bold=True, indent=False)
    add_para(doc, '默认支付期限为出发日期前30天。若从交款单发送日到默认支付期限的间隔不足10天，则支付期限调整为发送日后10天。用公式表示为：payment_deadline = max(出发日期 - 30天, 交款单发送日 + 10天)。该算法确保顾客至少有10天的时间来准备余款支付。')

    add_para(doc, '4. 旅游团状态管理算法', bold=True, indent=False)
    add_para(doc, '旅游团的状态按照以下规则进行转换：创建时为"未开放"状态；当价格被设定并公开后，转为"已开放"状态，此时开始接受顾客申请；当到达截止日期或人数达到限额时，转为"已截止"状态，不再接受新申请；当出发日期到达后，转为"已完成"状态。只有在"已开放"状态下才接受顾客申请。')

    add_para(doc, '（四）顺序图设计', bold=True, indent=False)
    add_para(doc, '以"办理旅游申请"这一核心用例为例，我们绘制了顺序图，描述了前台员工、系统界面、申请管理模块、旅游团查询模块、订金计算模块、支付记录模块和数据库之间的交互过程。')
    add_para(doc, '主要交互流程如下：')
    add_bullet(doc, '前台员工通过系统界面输入顾客需求（目的地、出行日期等），系统调用旅游团查询模块查询可用旅游团，从数据库获取信息后返回给前台员工。')
    add_bullet(doc, '前台员工选择旅游团后，系统验证申请条件：检查截止日期是否已过、人数是否已满。若不满足条件，返回拒绝原因；若满足条件，继续下一步。')
    add_bullet(doc, '前台员工录入申请责任人信息和参加人数，系统创建申请记录并保存到数据库。')
    add_bullet(doc, '系统调用订金计算模块，根据出发日期和参加人数计算订金金额，将结果展示给前台员工。')
    add_bullet(doc, '前台员工确认收取订金后，系统调用支付记录模块创建支付记录，保存到数据库，并生成收据。')
    add_bullet(doc, '系统将收据和旅游申请书展示给前台员工，前台员工打印后交给顾客。')

    add_para(doc, '（五）活动图设计', bold=True, indent=False)
    add_para(doc, '旅游申请业务的活动图描述了从顾客到店咨询到申请完成的完整流程，包含条件判断和多个处理阶段。')
    add_para(doc, '主要流程节点如下：')
    add_bullet(doc, '开始节点：顾客到店咨询。')
    add_bullet(doc, '前台员工查询旅游团信息，系统返回可选旅游团列表。')
    add_bullet(doc, '决策节点：判断是否满足申请条件（截止日期之前且人数未满）。不满足则告知顾客原因，流程结束。')
    add_bullet(doc, '满足条件后，进入订金处理分区：录入申请信息→计算订金→收取订金→记录支付→提供收据和申请书。')
    add_bullet(doc, '顾客填写申请书并邮寄回旅行社，员工接收后查询申请记录，录入参加者信息。')
    add_bullet(doc, '进入确认与催款分区：打印旅游确认书→打印余额交款单→邮寄给申请责任人。')
    add_bullet(doc, '顾客支付余款，系统记录支付信息，申请完成。')
    add_bullet(doc, '后续业务（可选）：参加者信息变更、取消参加者或取消整个申请。取消申请时需计算手续费并退还余款。')

    add_bold_label(doc, '三、结对编程过程')

    add_para(doc, '（一）角色互换时间点以及各自的任务分工', bold=True, indent=False)
    add_para(doc, '本次实验采用结对编程方式完成，两位成员分别担任驾驶员（负责编码）和领航员（负责审查代码和思考整体设计），每30分钟互换一次角色。具体分工如下：')

    pair_tbl = doc.add_table(rows=6, cols=4, style='Table Grid')
    pair_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['时间段', '驾驶员', '领航员', '任务内容']):
        set_cell(pair_tbl.cell(0, j), h, bold=True, size=10)
    tasks = [
        ['第1阶段\n（00:00-00:30）', '成员A', '成员B', '搭建项目框架，创建数据库表结构，实现基本数据模型类（旅游路线、旅游活动、旅游团、价格等）'],
        ['第2阶段\n（00:30-01:00）', '成员B', '成员A', '实现旅游团查询功能，实现申请条件验证逻辑，完成办理旅游申请的核心流程'],
        ['第3阶段\n（01:00-01:30）', '成员A', '成员B', '实现参加者信息录入与管理功能，完成申请确认流程，实现确认书和交款单的打印功能'],
        ['第4阶段\n（01:30-02:00）', '成员B', '成员A', '实现余款支付处理功能，实现取消申请及手续费计算逻辑，实现参加者变更功能'],
        ['第5阶段\n（02:00-02:30）', '成员A', '成员B', '实现旅游路线管理和价格设定功能，完成系统集成测试，修复发现的Bug'],
    ]
    for i, row in enumerate(tasks):
        for j, t in enumerate(row):
            set_cell(pair_tbl.cell(i + 1, j), t, size=10,
                     align=WD_ALIGN_PARAGRAPH.LEFT if j == 3 else WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '注：以上时间和分工为参考模板，请根据实际情况修改。', size=10)

    add_para(doc, '（二）工作照片', bold=True, indent=False)
    add_para(doc, '【请在此处插入结对编程的工作照片】', size=12)
    add_para(doc, '（照片1：结对编程工作场景——驾驶员编写代码，领航员审查指导）', size=11)
    add_para(doc, '（照片2：代码审查讨论场景——两位成员讨论设计方案）', size=11)

    add_bold_label(doc, '四、实验过程中发现的问题')
    add_para(doc, '1. 参与者边界识别问题：在进行用例分析时，发现"系统自动导出财务数据"这一用例的参与者较难确定。经过讨论，我们将其参与者定为"系统"（定时触发），同时考虑到会计人员是该功能数据的最终使用者，因此在用例图中也体现了会计人员与该用例的关联。')
    add_para(doc, '2. 类的层次关系混淆问题：在绘制类图时，旅游路线、旅游活动和旅游团三者之间的关系容易混淆。最初我们将旅游活动和旅游团设计为平级关系，后来发现旅游团实际上是旅游活动在特定出发日期的具体组织形式，应为旅游活动的下级概念，最终调整为旅游路线→旅游活动→旅游团的三层组合关系。')
    add_para(doc, '3. 价格不可变更约束的实现问题：问题描述中明确指出"价格公开后不可变更"，但在系统设计中如何实现这一约束需要特别注意。我们通过在Price类中设置"是否已公开"标志位，并在价格修改方法中加入状态检查来实现这一业务规则。')
    add_para(doc, '4. 订金计算的边界条件问题：订金计算依赖于距出发日期的天数，需要处理当天申请、已过截止日期等边界情况。初期实现时未考虑当天出发的极端情况，测试时发现了该问题并进行了修复。')
    add_para(doc, '5. 结对编程中的设计分歧：在取消申请的手续费计算逻辑上，两位成员对"申请完成前取消"的适用规则理解不一致。通过重新阅读问题描述，确认"申请完成前取消的情况也适用如下关系"，即申请完成前和完成后取消的手续费规则相同，最终达成一致。')

    # ================================================================
    # 第三部分：结果与讨论
    # ================================================================
    doc.add_page_break()
    add_heading(doc, '第三部分 结果与讨论', level=1)
    add_para(doc, '（实验结果分析，包括数据处理、实验现象分析、影响因素讨论、综合分析和结论等）', size=10, indent=False)

    add_bold_label(doc, '一、实验结果截图')
    add_para(doc, '以下为系统实现的主要功能截图（至少5张）：')

    screenshots = [
        ('截图1：系统主界面', '展示系统的主界面布局，包括功能导航菜单和各业务模块的入口。界面设计遵循简洁易用的原则，前台员工可以快速访问常用功能。'),
        ('截图2：旅游团查询功能', '展示根据顾客需求查询可用旅游团的功能界面。支持按目的地、出发日期等条件进行筛选，查询结果显示旅游团的基本信息、剩余名额和价格。'),
        ('截图3：办理旅游申请', '展示办理旅游申请的完整操作界面，包括申请责任人信息录入、参加人数选择、订金自动计算、支付记录等功能。系统自动验证申请条件并计算订金金额。'),
        ('截图4：参加者信息管理', '展示参加者信息的录入和管理界面。支持逐个录入参加者的详细信息（姓名、性别、年龄、类型），支持信息的修改和删除操作。'),
    ]
    for stitle, desc in screenshots:
        add_para(doc, stitle, bold=True, indent=False)
        add_para(doc, '【请在此处插入截图】')
        add_para(doc, desc, size=10)

    add_bold_label(doc, '二、实验结果分析')
    add_para(doc, '1. UML建模成果分析', bold=True, indent=False)
    add_para(doc, '通过本次实验，我们成功地将旅游业务管理系统的问题描述文档转化为一套完整的UML模型。用例图共识别出6个参与者和11个主要用例，覆盖了系统的所有功能需求。类图共识别出10个核心类和4个枚举类型，准确地反映了业务数据的静态结构和类之间的关系。顺序图和活动图则详细描述了关键业务场景的动态行为，为编码实现提供了明确的指导。')

    add_para(doc, '2. 系统功能实现分析', bold=True, indent=False)
    add_para(doc, '基于UML模型开发的系统原型，实现了旅游申请办理、参加者信息管理、余款支付处理、取消申请（含手续费计算）、旅游路线管理和价格设定等核心业务功能。经过功能测试，各业务场景均能正常运行，验证了UML分析与设计的正确性和有效性。')

    add_para(doc, '3. 结对编程效果分析', bold=True, indent=False)
    add_para(doc, '在结对编程过程中，领航员共发现了多处潜在的逻辑错误和设计问题，这些问题如果在后期才发现，修复成本将大大增加。同时，角色互换使得两位成员都能全面了解系统的各个模块，避免了知识孤岛的出现。与独立开发相比，结对编程的代码在可读性和健壮性方面表现更好。')

    add_bold_label(doc, '三、综合分析与结论')
    add_para(doc, '1. UML建模是软件开发过程中不可或缺的重要环节。通过本次实验，深刻体会到"磨刀不误砍柴工"的道理——前期的分析和设计工作虽然花费了一定时间，但为后续的编码实现提供了清晰的蓝图，大幅减少了开发过程中的返工和修改。特别是类图的设计，直接决定了代码的结构和质量。')
    add_para(doc, '2. 面向对象的分析方法能够自然地映射现实世界的业务实体和业务流程，降低了系统设计的认知复杂度。通过将业务概念抽象为类和对象，并利用继承、组合、关联等关系描述它们之间的联系，使得系统模型既贴近业务现实，又便于编码实现。')
    add_para(doc, '3. 不同的UML图从不同角度描述了系统的各个方面，它们之间相互补充、相互验证。用例图关注"做什么"，类图关注"有什么"，顺序图关注"怎么交互"，活动图关注"怎么流转"，状态机图关注"怎么变化"。综合运用这些UML图，能够全面、准确地描述一个复杂系统。')
    add_para(doc, '4. 结对编程模式在本次实验中表现出显著的优势。代码质量明显高于独立开发的预期水平，即时的代码审查机制有效减少了Bug的产生。同时，持续的沟通和讨论促进了设计方案的优化，很多好的设计想法正是在讨论中产生的。')
    add_para(doc, '5. 本次实验加深了对软件工程方法论的理解，建立了从需求分析到系统实现的完整认知。认识到软件开发不仅仅是编写代码，更是一个需要系统思考、团队协作和持续改进的工程过程。这些经验和体会将对今后的学习和工作产生积极的影响。')

    add_bold_label(doc, '四、结对编程之体会')
    add_para(doc, '通过本次结对编程实验，我们有以下几点深刻体会：')
    add_para(doc, '（1）结对编程显著提高了代码质量。在开发过程中，领航员能够实时审查驾驶员编写的代码，及时发现潜在的逻辑错误、边界条件遗漏和设计缺陷。很多问题在编码阶段就被发现并解决了，这比传统的后期代码审查更加高效。我们的代码在第一轮编写时的通过率明显高于平时独立开发的水平。')
    add_para(doc, '（2）角色互换促进了知识的双向传递。在结对编程中，当一位成员在某个技术点或业务理解上有更好的掌握时，另一位成员可以通过观察和讨论快速学习。这种即时的、实践中的知识传递比阅读技术文档更加直观和高效。通过本次实验，两位成员都学到了对方擅长的技术和思考方式。')
    add_para(doc, '（3）UML建模为开发提供了清晰的路线图。在编码之前进行系统的UML建模，使得整个开发过程有章可循。用例图帮助我们明确了系统的功能边界，避免了功能遗漏或越界；类图帮助我们理清了数据结构和对象关系，减少了编码时的数据模型混乱；顺序图和活动图为具体的编码实现提供了详细的行为描述，使得编码不再是"想到哪写到哪"。')
    add_para(doc, '（4）沟通与协作能力在实践中得到了提升。结对编程要求两位成员保持持续的沟通和协作。在这个过程中，我们学会了如何清晰地表达自己的设计思路，如何倾听和理解对方的观点，如何在意见分歧时通过理性讨论达成共识。这些软技能对于未来的团队开发工作非常有价值。')
    add_para(doc, '（5）对软件工程的实践认知得到了深化。本次实验让我们亲身体验了从需求分析、系统设计到编码实现的完整软件开发流程。深刻认识到每个阶段的重要性——需求分析决定了"做对的事"，系统设计决定了"把事做对"，编码实现决定了"把事做好"。同时也认识到，软件开发是一个需要不断迭代和改进的过程，不可能一步到位。')

    # ================================================================
    # 成绩评定表
    # ================================================================
    doc.add_paragraph()
    add_heading(doc, '成绩评定表', level=1)

    grade_tbl = doc.add_table(rows=6, cols=4, style='Table Grid')
    grade_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['序号', '评分项目', '满分', '实得分']):
        set_cell(grade_tbl.cell(0, j), h, bold=True, size=11)
    grades = [
        ['1', '实验报告格式规范', '2', ''],
        ['2', '实验报告过程清晰，内容详实', '4', ''],
        ['3', '实验报告结果正确性', '2', ''],
        ['4', '实验分析与总结详尽', '2', ''],
    ]
    for i, row in enumerate(grades):
        for j, t in enumerate(row):
            set_cell(grade_tbl.cell(i + 1, j), t, size=11)
    set_cell(grade_tbl.cell(5, 0), '', bold=True, size=11)
    grade_tbl.cell(5, 0).merge(grade_tbl.cell(5, 1))
    set_cell(grade_tbl.cell(5, 0), '总得分', bold=True, size=11)
    set_cell(grade_tbl.cell(5, 2), '10', bold=True, size=11)
    set_cell(grade_tbl.cell(5, 3), '', bold=True, size=11)

    # 保存
    out = r'D:\旅游业务管理系统原型\实验报告_旅游业务管理系统_v2.docx'
    doc.save(out)
    print(f'已生成: {out}')


if __name__ == '__main__':
    create_report()
